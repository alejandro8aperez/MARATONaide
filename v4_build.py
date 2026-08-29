# -*- coding: utf-8 -*-
"""
Genera la diagramación VERSIÓN 4 de MARATONaide a partir de 'Maratonaide ver 3.docx'.

Toma el contenido textual de la v3 y lo re-diagrama con estilos profesionales de
libro listos para publicación en este orden:
  Portada -> Página legal -> Dedicatoria -> Contraportada (sinopsis) -> Índice
  -> Prólogo -> 23 capítulos -> Nota del autor -> FIN
"""

import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:\Postgres\MARATONaide\Maratonaide ver 3.docx"
OUT = r"C:\Postgres\MARATONaide\MARATONAIDE ver 4.docx"

GOLD = RGBColor(0xC9, 0x9A, 0x2C)
DARK = RGBColor(0x20, 0x20, 0x24)
GREY = RGBColor(0x5A, 0x5A, 0x60)


def set_run(run, size=None, bold=None, italic=None, color=None, name="Georgia"):
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    rf.set(qn('w:eastAsia'), name)


def add_bottom_border(paragraph, color='C99A2C', sz='8'):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def detect_section(text):
    up = text.strip().upper().replace('\u00c1','A').replace('\u00c9','E').replace('\u00cd','I').replace('\u00d3','O').replace('\u00da','U').replace('\u00d1','N')
    if up.startswith('CAP') and ('KM' in up):
        return 'capitulo'
    if 'DEDICATORIA' in up:
        return 'dedicatoria'
    if 'SINOPSIS' in up or 'CONTRAPORTADA' in up:
        return 'sinopsis'
    if 'PROLOGO' in up:
        return 'prologo'
    if 'NOTA DEL AUTOR' in up:
        return 'nota'
    if 'PAGINA LEGAL' in up or 'LEGAL' in up or 'CREDITOS' in up:
        return 'legal'
    if 'MARATONAIDE' in up and len(up) < 20:
        return 'titulo'
    return 'otro'


def load_source():
    d = docx.Document(SRC)
    paras = []
    for p in d.paragraphs:
        paras.append((p.text, p.style.name if p.style else 'Normal', p.alignment))
    return paras


def main():
    paras = load_source()

    # Estructurar en secciones por Heading 1
    sections = []
    current_kind = None
    current_buf = []
    for text, style, align in paras:
        if style == 'Heading 1':
            if current_kind is not None:
                sections.append((current_kind, current_buf))
            current_kind = detect_section(text)
            current_buf = [(text, style, align)]
        else:
            current_buf.append((text, style, align))
    if current_kind is not None:
        sections.append((current_kind, current_buf))

    # Separar por tipo, conservando capital del capítulo
    legal = ded = syn = prologue = nota = None
    capitulos = []
    for kind, buf in sections:
        if kind == 'legal':
            legal = buf
        elif kind == 'dedicatoria':
            ded = buf
        elif kind == 'sinopsis':
            syn = buf
        elif kind == 'prologo':
            prologue = buf
        elif kind == 'nota':
            nota = buf
        elif kind == 'capitulo':
            capitulos.append(buf)

    doc = Document()

    sec = doc.sections[0]
    sec.page_width = Inches(6)
    sec.page_height = Inches(9)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    sec.top_margin = Inches(0.9)
    sec.bottom_margin = Inches(0.9)

    base = doc.styles['Normal']
    base.font.name = 'Georgia'
    base.font.size = Pt(11)
    base.font.color.rgb = DARK
    base.paragraph_format.line_spacing = 1.35
    base.paragraph_format.space_after = Pt(9)
    rpr = base.element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    rf.set(qn('w:eastAsia'), 'Georgia')

    # ---------- helpers ----------
    def blank(n=1):
        for _ in range(n):
            doc.add_paragraph()

    def page_title(title, small=None):
        if small:
            s = doc.add_paragraph()
            s.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run(s.add_run(small), size=12, italic=True, color=GOLD)
            s.paragraph_format.space_before = Pt(6)
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(h.add_run(title), size=19, bold=True, color=DARK)
        add_bottom_border(h)
        h.paragraph_format.space_after = Pt(18)

    def body(text, centered=False, no_indent=False, italic=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.JUSTIFY
        if not centered and not no_indent:
            p.paragraph_format.first_line_indent = Inches(0.3)
        set_run(p.add_run(text), size=11, italic=italic)
        return p

    # ============ 1. PORTADA ============
    blank(4)
    ep = doc.add_paragraph(); ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(ep.add_run("\u201cUno no corre pa' llegar. Uno corre pa' dejar de huir.\u201d"),
            size=13, italic=True, color=GREY)
    blank(1)
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(t.add_run("MARATONaide"), size=42, bold=True, color=GOLD)
    blank(2)
    st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(st.add_run("42 km \u00b7 22 etapas \u00b7 195 metros"),
            size=16, italic=True, color=DARK)
    blank(3)
    au = doc.add_paragraph(); au.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(au.add_run("ALEJANDRO OCHOA P\u00c9REZ"), size=14, bold=True, color=DARK)
    ver = doc.add_paragraph(); ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(ver.add_run("Versi\u00f3n 4 \u00b7 Diagramaci\u00f3n para publicaci\u00f3n"),
            size=9, italic=True, color=GREY)
    doc.add_page_break()

    # ============ 2. PÁGINA LEGAL ============
    if legal:
        for text, style, align in legal:
            tt = text.strip()
            if not tt:
                continue
            up = tt.upper()
            if 'MARATONAIDE' == up or tt.upper().startswith('MARATONAIDE') and len(tt) < 20:
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run(p.add_run(tt), size=12, bold=True, color=GOLD)
                p.paragraph_format.space_before = Pt(0)
            elif 'PAGINA LEGAL' in up or 'LEGAL' in up or 'CREDITOS' in up:
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run(p.add_run(tt), size=14, bold=True, color=DARK)
                p.paragraph_format.space_before = Pt(160)
            else:
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run(p.add_run(tt), size=9, color=GREY)
    doc.add_page_break()

    # ============ 3. DEDICATORIA ============
    if ded:
        blank(6)
        for text, style, align in ded:
            tt = text.strip()
            if not tt or tt.upper() == 'DEDICATORIA':
                continue
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run(p.add_run(tt), size=13, italic=True, color=DARK)
    doc.add_page_break()

    # ============ 4. CONTRAPORTADA / SINOPSIS ============
    if syn:
        page_title("Contraportada", "Sinopsis")
        for text, style, align in syn:
            tt = text.strip()
            if not tt or tt.upper().startswith('CONTRAPORTADA') or tt.upper().startswith('SINOPSIS'):
                continue
            body(tt)
    doc.add_page_break()

    # ============ 5. ÍNDICE ============
    page_title("\u00cdndice")
    # Nota: incluimos portada/legal/dedicatoria/sinopsis/prólogo/nota si existe
    toc_items = []
    if prologue:
        toc_items.append("Pr\u00f3logo")
    for i, buf in enumerate(capitulos, 1):
        toc_items.append(buf[0][0].strip())
    if nota:
        toc_items.append("Nota del autor")
    for item in toc_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_run(p.add_run(item), size=11, color=DARK)
        p.paragraph_format.space_after = Pt(6)
    doc.add_page_break()

    # ============ 6. PRÓLOGO ============
    if prologue:
        first = prologue[0][0].strip()
        page_title(first)
        for text, style, align in prologue[1:]:
            tt = text.strip()
            if not tt:
                continue
            if tt.isupper() and len(tt) < 90:
                s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run(s.add_run(tt), size=12, bold=True, color=DARK)
                s.paragraph_format.space_before = Pt(8)
            else:
                body(tt)
        doc.add_page_break()

    # ============ 7. CAPÍTULOS ============
    total = len(capitulos)
    for idx, buf in enumerate(capitulos):
        first = buf[0][0].strip()
        page_title(first)
        for text, style, align in buf[1:]:
            tt = text.strip()
            if not tt:
                continue
            if tt.isupper() and len(tt) < 90:
                s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run(s.add_run(tt), size=12, bold=True, color=DARK)
                s.paragraph_format.space_before = Pt(8)
            else:
                body(tt)
        if idx != total - 1:
            doc.add_page_break()

    # ============ 8. NOTA DEL AUTOR ============
    if nota:
        doc.add_page_break()
        first = nota[0][0].strip()
        page_title(first)
        for text, style, align in nota[1:]:
            tt = text.strip()
            if not tt:
                continue
            if tt.isupper() and len(tt) < 90:
                s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run(s.add_run(tt), size=12, bold=True, color=DARK)
                s.paragraph_format.space_before = Pt(8)
            else:
                body(tt)

    # FIN
    doc.add_page_break()
    fin = doc.add_paragraph(); fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fin.paragraph_format.space_before = Pt(160)
    set_run(fin.add_run("FIN"), size=24, bold=True, color=GOLD)

    doc.save(OUT)
    print("Guardado:", OUT)
    print("Secciones:", len(sections))
    print("Capítulos:", len(capitulos))
    for b in capitulos:
        print("  -", b[0][0].strip())
    print("Prólogo:", bool(prologue), "| Legal:", bool(legal),
          "| Dedicatoria:", bool(ded), "| Sinopsis:", bool(syn), "| Nota:", bool(nota))


if __name__ == "__main__":
    main()

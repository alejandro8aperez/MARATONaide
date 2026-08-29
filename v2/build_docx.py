# -*- coding: utf-8 -*-
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:\Postgres\MARATONaide\MARATONaide - Version 2.0.md"
DST = r"C:\Postgres\MARATONaide\MARATONaide - Version 2.0.docx"

with open(SRC, encoding="utf-8") as f:
    md = f.read()

blocks = [b.strip() for b in re.split(r"\n\s*\n", md) if b.strip()]

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
pf = style.paragraph_format
pf.line_spacing = 1.5

TOKEN = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")

def add_para(text, center=False, size=12, bold=False, italic=False, page_break=False):
    p = doc.add_paragraph()
    if page_break:
        pPr = p._p.get_or_add_pPr()
        pPr.append(OxmlElement("w:pageBreakBefore"))
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    for part in TOKEN.split(text):
        if not part:
            continue
        b, i = bold, italic
        t = part
        if part.startswith("**") and part.endswith("**"):
            b, t = True, part[2:-2]
            if t.startswith("*") and t.endswith("*"):
                i, t = True, t[1:-1]
        elif part.startswith("*") and part.endswith("*"):
            i, t = True, part[1:-1]
        r = p.add_run(t.replace("**","").replace("*",""))
        r.font.name = "Times New Roman"; r.font.size = Pt(size); r.bold = b; r.italic = i
    return p

first_h1 = True
for b in blocks:
    first_line = b.split("\n")[0]
    m = re.match(r"^(#{1,6})\s+(.*)$", first_line)
    if m and all(not l.strip().startswith(("#","*")) or l is first_line for l in [first_line]):
        level = len(m.group(1)); text = m.group(2).strip()
        if level == 1:
            add_para(text, center=True, size=18, bold=True, page_break=not first_h1)
            first_h1 = False
        else:
            add_para(text, center=True, size=14, bold=True, page_break=True)
    else:
        for line in b.split("\n"):
            t = line.strip()
            if not t: continue
            full_bold_italic = t.startswith("**") and t.endswith("**") or (t.startswith("*") and t.endswith("*"))
            add_para(t)

doc.save(DST)

d2 = Document(DST)
print("parrafos:", len(d2.paragraphs))
print("palabras:", sum(len(p.text.split()) for p in d2.paragraphs))
leftover = [p.text[:60] for p in d2.paragraphs if "*" in p.text]
print("marcadores sobrantes:", len(leftover))
for x in leftover[:5]: print("  ", x)
heads = sum(1 for p in d2.paragraphs if p.runs and p.runs[0].bold and p.alignment == WD_ALIGN_PARAGRAPH.CENTER)
print("titulos centrados:", heads)